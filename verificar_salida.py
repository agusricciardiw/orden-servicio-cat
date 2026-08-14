# -*- coding: utf-8 -*-
"""
CONTROL FINAL antes de enviar.

Compara las ordenes generadas contra la planilla, SIN usar el generador:
lee el Excel por su cuenta y lee los Word ya escritos. Si las dos lecturas
coinciden, lo que se manda es lo que se cargo y valido.

    python verificar_salida.py
    python verificar_salida.py --xlsx "ruta\\al\\archivo.xlsx"

Chequea:
  1. Que cada servicio validado aparezca en el documento, y solo una vez
  2. Que no haya servicios de mas
  3. Que la suma de agentes por turno coincida
  4. Que el indice liste los mismos anexos que trae el documento
  5. Que ningun link del indice quede roto
"""
import sys, io, re, argparse, unicodedata
from pathlib import Path
from collections import Counter

import openpyxl
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
AQUI = Path(__file__).parent

ORDENES = [
    ('SEMANA', 'SEMANA', AQUI / 'salida' / 'Orden de servicio SEMANA.docx'),
    ('FINDE',  'FINDE',  AQUI / 'salida' / 'Orden de servicio FINDE.docx'),
]


def norm(s):
    """Normaliza para comparar: sin acentos, sin espacios de mas, mayusculas."""
    s = str(s or '').strip()
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(c for c in s if not unicodedata.combining(c))
    return re.sub(r'\s+', ' ', s).upper()


def clave(funcion, servicio):
    """Huella de un servicio: sirve para contar sin depender del formateo.
    Se saca el prefijo MISION, que el generador agrega en la tabla cuando el
    TIPO es MISION y por lo tanto no esta en la columna SERVICIO."""
    s = re.sub(r'^MISI[OÓ]N\s+', '', norm(servicio))
    return (s[:40], norm(funcion)[:70])


# ---------------------------------------------------------------- el Excel
def leer_planilla(xlsx, hoja):
    ws = openpyxl.load_workbook(xlsx, data_only=True)[hoja]
    cols, dot, val = {}, [], None
    for c in range(1, ws.max_column + 1):
        v = ws.cell(row=6, column=c).value
        if v is None:
            continue
        n = norm(v)
        if n in ('ID', 'SERVICIO', 'BASE', 'FUNCION'):
            cols.setdefault(n, c)
        elif n.startswith('AT '):
            dot.append(c)
        elif n == 'ANADIDO':
            val = c
    if val is None or not cols.get('ID'):
        sys.exit(f"No pude ubicar las columnas en la hoja {hoja}")

    servicios, agentes, sin_base = [], 0, 0
    for r in range(7, ws.max_row + 1):
        if not str(ws.cell(row=r, column=cols['ID']).value or '').strip():
            continue
        if ws.cell(row=r, column=val).value is not True:
            continue
        base = str(ws.cell(row=r, column=cols['BASE']).value or '').strip()
        if not base or base == '0':
            sin_base += 1
            continue
        servicios.append(clave(ws.cell(row=r, column=cols['FUNCION']).value,
                               ws.cell(row=r, column=cols['SERVICIO']).value))
        for c in dot:
            v = ws.cell(row=r, column=c).value
            agentes += int(v) if isinstance(v, (int, float)) else 0
    return servicios, agentes, sin_base


# ------------------------------------------------------------------ el Word
def leer_documento(ruta):
    doc = Document(ruta)
    servicios, agentes = [], 0
    anexos = []
    for t in doc.tables:
        filas = t.rows
        if len(filas) < 3:
            continue
        titulo = filas[0].cells[0].text.strip()
        cab = [norm(c.text) for c in filas[1].cells]
        if 'DIA' not in cab:                       # no es tabla de anexo
            continue

        # Ubicar las columnas por su encabezado, no por posicion
        def idx(nombre):
            return cab.index(nombre) if nombre in cab else None
        i_tarea, i_desc = idx('TAREA'), idx('DESCRIPCION')
        i_hora, i_base = idx('HORA'), idx('BASE')
        if None in (i_tarea, i_desc, i_hora, i_base):
            print(f"  !! No pude ubicar las columnas en el anexo {titulo!r}")
            continue

        anexos.append(titulo)
        for fila in filas[3:]:
            celdas = [c.text for c in fila.cells]
            if len(celdas) <= i_base:
                continue
            servicios.append(clave(celdas[i_desc], celdas[i_tarea]))
            for c in celdas[i_hora + 1:i_base]:    # dotacion por turno
                c = c.strip()
                if c.isdigit():
                    agentes += int(c)

    body = doc.element.body
    anclas = [h.get(W + 'anchor') for h in body.findall(f'.//{W}hyperlink')
              if h.get(W + 'anchor')]
    marcas = {b.get(W + 'name') for b in body.findall(f'.//{W}bookmarkStart')}
    indice = [c.text.split('\n')[0].strip()
              for t in doc.tables[:1] for fila in t.rows for c in fila.cells
              if c.text.strip().upper().startswith('ANEXO ')]
    return servicios, agentes, anexos, anclas, marcas, indice


def comparar(nombre, hoja, docx, xlsx):
    print('=' * 70)
    print(f"{nombre}   {docx.name}")
    print('=' * 70)
    if not docx.exists():
        print("  !! NO EXISTE el documento generado")
        return False

    esp, ag_esp, sin_base = leer_planilla(xlsx, hoja)
    obt, ag_obt, anexos, anclas, marcas, indice = leer_documento(docx)

    ok = True
    print(f"  servicios validados con base  : {len(esp)}")
    print(f"  servicios en el documento     : {len(obt)}")
    if sin_base:
        print(f"  (mas {sin_base} validado/s SIN BASE, que no entran a ningun anexo)")

    ce, co = Counter(esp), Counter(obt)
    faltan = ce - co
    sobran = co - ce
    if faltan:
        ok = False
        print(f"\n  !! FALTAN {sum(faltan.values())} servicio(s) en el documento:")
        for (s, f), n in list(faltan.items())[:12]:
            print(f"     x{n}  {s[:34]} | {f[:44]}")
    if sobran:
        ok = False
        print(f"\n  !! SOBRAN {sum(sobran.values())} servicio(s) en el documento:")
        for (s, f), n in list(sobran.items())[:12]:
            print(f"     x{n}  {s[:34]} | {f[:44]}")
    if not faltan and not sobran:
        print("  OK  cada servicio validado aparece una sola vez")

    print(f"\n  agentes sumados en la planilla: {ag_esp}")
    print(f"  agentes sumados en el documento: {ag_obt}")
    if ag_esp != ag_obt:
        ok = False
        print(f"  !! NO COINCIDE (diferencia {ag_obt - ag_esp})")
    else:
        print("  OK  la dotacion total coincide")

    print(f"\n  anexos en el documento: {len(anexos)}")
    for a in anexos:
        print(f"     - {a}")
    en_indice = [i.replace('ANEXO ', '') for i in indice]
    faltan_idx = [a for a in anexos if not any(norm(a) in norm(i) for i in en_indice)]
    if faltan_idx:
        ok = False
        print(f"  !! Estos anexos no figuran en el indice: {faltan_idx}")
    else:
        print(f"  OK  el indice lista los {len(en_indice)} anexos")

    rotos = [a for a in anclas if a not in marcas]
    if rotos:
        ok = False
        print(f"  !! links rotos: {rotos}")
    else:
        print(f"  OK  los {len(anclas)} links del indice tienen destino")
    return ok


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--xlsx', default=None)
    a = ap.parse_args()

    if a.xlsx:
        xlsx = Path(a.xlsx)
    else:
        cand = sorted((Path.home() / 'Downloads').glob('SISTEMA DE PREORDEN*.xlsx'),
                      key=lambda p: p.stat().st_mtime, reverse=True)
        if not cand:
            sys.exit("No encuentro la planilla en Descargas")
        xlsx = cand[0]
    print(f"\nPlanilla: {xlsx.name}\n")

    todo_ok = True
    for nombre, hoja, docx in ORDENES:
        if not comparar(nombre, hoja, docx, xlsx):
            todo_ok = False
        print()

    print('=' * 70)
    print("  TODO OK: lo generado coincide con la planilla" if todo_ok
          else "  HAY DIFERENCIAS: revisar antes de enviar")
    print('=' * 70)
    return 0 if todo_ok else 1


if __name__ == '__main__':
    sys.exit(main())
